from openpyxl import load_workbook,Workbook
import datetime
import os.path
from pprint import pprint
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dateutil.relativedelta import relativedelta

month_names = (
	u'января',
	u'февраля',
	u'марта',
	u'апреля',
	u'мая',
	u'июня',
	u'июля',
	u'августа',
	u'сентября',
	u'октября',
	u'ноября',
	u'декабря'
	)

def get_data_xls(filename):

	wb = load_workbook(filename=filename)
	ws = wb.active
	entries = []

	for row in ws.rows:
		entry = []
		for cell in row:
			entry.append(cell.value if cell.value else '')
		entries.append(entry)

	# check if last rows is None
	for entry in reversed(entries):
		if any(entry):
			break
		else:
			del entries[-1]

	# check if last columns is None
	for _ in entries[0]:
		last_column = [_ for entry in entries if entry[-1]]
		if last_column:
			break
		else:
			for index,_ in enumerate(entries):
				del entries[index][-1]

	return entries

def save_xls_data(filename, entries):


	if os.path.isfile(filename):
		wb = load_workbook(filename)
	else:
		wb = Workbook()
	ws = wb.active
	ws.title = u'База'

	for row_index,entry in enumerate(entries):
		for col_index,value in enumerate(entry):
			ws.cell(row=row_index+2,column=col_index+1).value = value

	wb.save(filename)

def get_doc_data(filename):

	document = Document(filename)
	lines = []
	for paragraph in document.paragraphs:
		line = paragraph.text
		if line != '':
			lines.append(line)

	for table in document.tables:
		for row in table.rows:
			for cell in row.cells:
				for paragraph in cell.paragraphs:
					line = paragraph.text
					if line != '':
						lines.append(line)

	return lines

def create_new_replaced_doc(fin,fout,replacements):

	document = Document(fin)

	for paragraph in document.paragraphs:
		for old_data,new_data in replacements.items():
			if old_data in paragraph.text:
				# Loop added to work with runs (strings with same style)
				for run in paragraph.runs:
					if old_data in run.text:
						run.text = run.text.replace(old_data, new_data)

	for table in document.tables:
		for row in table.rows:
			for cell in row.cells:
				for paragraph in cell.paragraphs:
					for old_data,new_data in replacements.items():
						if old_data in paragraph.text:
							# Loop added to work with runs (strings with same style)
							for run in paragraph.runs:
								if old_data in run.text:
									run.text = run.text.replace(old_data, new_data)

	document.save(fout)

def get_current_date_russian():

	now = datetime.datetime.now()
	return u'« %s » %s %sг.' % (
		now.day,
		month_names[now.month-1],
		now.year
		)

def get_truncated_line(line,length):

	if len(line)<6:
		return line

	if len(line)<=length:
		return line
	
	return line[:length//2 - 3] + ' ... ' + line[-length//2:]

def add_date(d,years=0,months=0):

	return d + relativedelta(years=+years,months=months)